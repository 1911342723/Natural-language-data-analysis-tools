"""
文件上传和解析API - 用于科学家团队模式
"""
import logging
import os
import json
from typing import Dict, Any, Optional
from fastapi import APIRouter, UploadFile, File, HTTPException
from pathlib import Path
import pandas as pd
import PyPDF2
from docx import Document

logger = logging.getLogger(__name__)

router = APIRouter(prefix="/api/team", tags=["team"])

# 支持的文件类型
ALLOWED_EXTENSIONS = {
    # 数据文件
    'csv': 'data',
    'xlsx': 'data',
    'xls': 'data',
    # 文档文件
    'pdf': 'document',
    'docx': 'document',
    'txt': 'document',
    'md': 'document',
    # 图片文件（暂不支持，未来可扩展）
    # 'jpg': 'image',
    # 'jpeg': 'image',
    # 'png': 'image',
}

MAX_FILE_SIZE = 50 * 1024 * 1024  # 50MB


def get_file_type(filename: str) -> Optional[str]:
    """获取文件类型"""
    ext = filename.rsplit('.', 1)[-1].lower() if '.' in filename else None
    return ALLOWED_EXTENSIONS.get(ext)


async def parse_csv(file_path: str) -> Dict[str, Any]:
    """解析CSV文件"""
    try:
        df = pd.read_csv(file_path, encoding='utf-8')
        
        return {
            'type': 'data',
            'format': 'csv',
            'rows': len(df),
            'columns': list(df.columns),
            'summary': {
                'shape': df.shape,
                'dtypes': df.dtypes.astype(str).to_dict(),
                'head': df.head(5).to_dict('records'),
                'describe': df.describe().to_dict() if len(df) > 0 else {}
            },
            'preview': df.head(10).to_string()
        }
    except UnicodeDecodeError:
        # 尝试GBK编码
        df = pd.read_csv(file_path, encoding='gbk')
        return await parse_csv(file_path)  # 递归调用
    except Exception as e:
        logger.error(f"CSV解析失败: {e}")
        raise HTTPException(status_code=400, detail=f"CSV解析失败: {str(e)}")


async def parse_excel(file_path: str) -> Dict[str, Any]:
    """解析Excel文件"""
    try:
        # 读取所有sheet
        excel_file = pd.ExcelFile(file_path)
        sheets = {}
        
        for sheet_name in excel_file.sheet_names:
            df = pd.read_excel(file_path, sheet_name=sheet_name)
            sheets[sheet_name] = {
                'rows': len(df),
                'columns': list(df.columns),
                'preview': df.head(5).to_dict('records')
            }
        
        # 默认使用第一个sheet
        df = pd.read_excel(file_path, sheet_name=0)
        
        return {
            'type': 'data',
            'format': 'excel',
            'sheets': list(excel_file.sheet_names),
            'rows': len(df),
            'columns': list(df.columns),
            'summary': {
                'shape': df.shape,
                'dtypes': df.dtypes.astype(str).to_dict(),
                'head': df.head(5).to_dict('records'),
                'describe': df.describe().to_dict() if len(df) > 0 else {}
            },
            'all_sheets': sheets,
            'preview': df.head(10).to_string()
        }
    except Exception as e:
        logger.error(f"Excel解析失败: {e}")
        raise HTTPException(status_code=400, detail=f"Excel解析失败: {str(e)}")


async def parse_pdf(file_path: str) -> Dict[str, Any]:
    """解析PDF文件"""
    try:
        with open(file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            num_pages = len(pdf_reader.pages)
            
            # 提取前几页文本
            text_content = []
            for page_num in range(min(5, num_pages)):  # 最多读取前5页
                page = pdf_reader.pages[page_num]
                text = page.extract_text()
                text_content.append(f"=== 第 {page_num + 1} 页 ===\n{text}")
            
            full_text = "\n\n".join(text_content)
            
            return {
                'type': 'document',
                'format': 'pdf',
                'pages': num_pages,
                'text': full_text,
                'preview': full_text[:1000] + ('...' if len(full_text) > 1000 else ''),
                'word_count': len(full_text.split())
            }
    except Exception as e:
        logger.error(f"PDF解析失败: {e}")
        raise HTTPException(status_code=400, detail=f"PDF解析失败: {str(e)}")


async def parse_docx(file_path: str) -> Dict[str, Any]:
    """解析Word文档"""
    try:
        doc = Document(file_path)
        
        # 提取所有段落
        paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
        full_text = "\n\n".join(paragraphs)
        
        # 提取表格
        tables = []
        for table in doc.tables:
            table_data = []
            for row in table.rows:
                row_data = [cell.text for cell in row.cells]
                table_data.append(row_data)
            tables.append(table_data)
        
        return {
            'type': 'document',
            'format': 'docx',
            'paragraphs': len(paragraphs),
            'tables': len(tables),
            'text': full_text,
            'preview': full_text[:1000] + ('...' if len(full_text) > 1000 else ''),
            'word_count': len(full_text.split())
        }
    except Exception as e:
        logger.error(f"Word解析失败: {e}")
        raise HTTPException(status_code=400, detail=f"Word解析失败: {str(e)}")


async def parse_text(file_path: str) -> Dict[str, Any]:
    """解析纯文本文件"""
    try:
        with open(file_path, 'r', encoding='utf-8') as f:
            text = f.read()
        
        return {
            'type': 'document',
            'format': 'txt',
            'text': text,
            'preview': text[:1000] + ('...' if len(text) > 1000 else ''),
            'lines': len(text.split('\n')),
            'word_count': len(text.split())
        }
    except UnicodeDecodeError:
        # 尝试GBK编码
        with open(file_path, 'r', encoding='gbk') as f:
            text = f.read()
        return await parse_text(file_path)  # 递归调用
    except Exception as e:
        logger.error(f"文本解析失败: {e}")
        raise HTTPException(status_code=400, detail=f"文本解析失败: {str(e)}")


@router.post("/upload_file")
async def upload_file(file: UploadFile = File(...)):
    """
    上传并解析文件
    
    支持的文件类型：
    - 数据文件：CSV, Excel (.xlsx, .xls)
    - 文档：PDF, Word (.docx), TXT, Markdown
    """
    try:
        # 检查文件类型
        file_type = get_file_type(file.filename)
        if not file_type:
            raise HTTPException(
                status_code=400,
                detail=f"不支持的文件类型。支持：{', '.join(ALLOWED_EXTENSIONS.keys())}"
            )
        
        # 检查文件大小
        file.file.seek(0, 2)  # 移到文件末尾
        file_size = file.file.tell()
        file.file.seek(0)  # 重置到开头
        
        if file_size > MAX_FILE_SIZE:
            raise HTTPException(
                status_code=400,
                detail=f"文件过大（{file_size / 1024 / 1024:.1f}MB），最大支持 {MAX_FILE_SIZE / 1024 / 1024}MB"
            )
        
        # 保存文件到临时目录
        upload_dir = Path("./uploads/team_files")
        upload_dir.mkdir(parents=True, exist_ok=True)
        
        file_path = upload_dir / file.filename
        with open(file_path, "wb") as f:
            content = await file.read()
            f.write(content)
        
        logger.info(f"📁 文件已上传: {file.filename} ({file_size / 1024:.1f}KB)")
        
        # 根据文件类型解析
        ext = file.filename.rsplit('.', 1)[-1].lower()
        
        if ext == 'csv':
            parsed_data = await parse_csv(str(file_path))
        elif ext in ['xlsx', 'xls']:
            parsed_data = await parse_excel(str(file_path))
        elif ext == 'pdf':
            parsed_data = await parse_pdf(str(file_path))
        elif ext == 'docx':
            parsed_data = await parse_docx(str(file_path))
        elif ext in ['txt', 'md']:
            parsed_data = await parse_text(str(file_path))
        else:
            raise HTTPException(status_code=400, detail="不支持的文件类型")
        
        # 添加通用信息
        parsed_data['filename'] = file.filename
        parsed_data['size'] = file_size
        parsed_data['file_path'] = str(file_path)
        
        logger.info(f"✅ 文件解析完成: {file.filename}, 类型: {file_type}")
        
        return {
            'success': True,
            'data': parsed_data
        }
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"文件上传失败: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail=f"文件上传失败: {str(e)}")

